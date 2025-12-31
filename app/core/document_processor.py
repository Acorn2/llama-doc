"""
通用文档处理器
支持PDF、TXT、DOC、DOCX文件的处理
"""
import logging
import os
from typing import List, Dict, Any, Optional
from pathlib import Path

from llama_index.core.schema import Document
from llama_index.readers.file import PyMuPDFReader
from llama_index.core import SimpleDirectoryReader
from llama_index.core.node_parser import SentenceSplitter

# 修改logger名称，确保与日志配置中一致
logger = logging.getLogger("app.core.document_processor")

class DocumentProcessor:
    """通用文档处理器，支持多种文件格式"""
    
    # 支持的文件类型及其MIME类型
    SUPPORTED_TYPES = {
        '.pdf': 'application/pdf',
        '.txt': 'text/plain',
        '.doc': 'application/msword',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }
    
    def __init__(self):
        # 延迟初始化，只在需要时创建pdf_reader
        self._pdf_reader = None
    
    @property
    def pdf_reader(self):
        """懒加载PDF阅读器"""
        if self._pdf_reader is None:
            from llama_index.readers.file import PyMuPDFReader
            self._pdf_reader = PyMuPDFReader()
            logger.info("PyMuPDFReader初始化成功")
        return self._pdf_reader
    
    def is_supported_file(self, filename: str) -> bool:
        """检查文件是否支持"""
        file_ext = Path(filename).suffix.lower()
        return file_ext in self.SUPPORTED_TYPES
    
    def get_content_type(self, filename: str) -> str:
        """获取文件的MIME类型"""
        file_ext = Path(filename).suffix.lower()
        return self.SUPPORTED_TYPES.get(file_ext, 'application/octet-stream')
    
    def load_data(self, file_path: str) -> List[Document]:
        """
        加载文档并返回Document列表
        
        Args:
            file_path: 文件路径
            
        Returns:
            List[Document]: LlamaIndex Document对象列表
        """
        file_ext = Path(file_path).suffix.lower()
        logger.info(f"处理文件: {file_path}, 类型: {file_ext}")
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            error_msg = f"文件不存在: {file_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
        
        # 检查文件是否为空
        if os.path.getsize(file_path) == 0:
            error_msg = f"文件为空: {file_path}"
            logger.error(error_msg)
            raise ValueError(error_msg)
        
        # 检查文件扩展名是否在支持列表中
        if file_ext not in self.SUPPORTED_TYPES:
            error_msg = f"文件扩展名不在支持列表中: {file_ext}, 支持的类型: {list(self.SUPPORTED_TYPES.keys())}"
            logger.warning(error_msg)
            # 根据文件路径尝试识别真实类型
            # 这里不直接报错，而是记录警告并继续尝试处理
            
        try:
            logger.info(f"正在处理文件: {file_path}, 扩展名: {file_ext}")
            
            if file_ext == '.pdf':
                try:
                    logger.info("开始加载PDF文档...")
                    return self._load_pdf(file_path)
                except Exception as e:
                    logger.error(f"PDF加载失败: {str(e)}")
                    raise ValueError(f"文件不是有效的PDF格式: {str(e)}")
            elif file_ext == '.txt':
                try:
                    logger.info("开始加载TXT文档...")
                    return self._load_txt(file_path)
                except Exception as e:
                    logger.error(f"TXT加载失败: {str(e)}")
                    raise ValueError(f"无法加载TXT文件: {str(e)}")
            elif file_ext in ['.doc', '.docx']:
                try:
                    logger.info(f"开始加载Word文档({file_ext})...")
                    return self._load_word(file_path)
                except ImportError as e:
                    logger.error(f"缺少Word处理库: {str(e)}")
                    raise
                except Exception as e:
                    logger.error(f"Word文档加载失败: {str(e)}")
                    raise ValueError(f"无法加载Word文档: {str(e)}")
            else:
                error_msg = f"不支持的文件类型: {file_ext}"
                logger.error(error_msg)
                raise ValueError(error_msg)
        except Exception as e:
            logger.error(f"加载文件失败: {str(e)}", exc_info=True)
            raise
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """加载PDF文件"""
        documents = self.pdf_reader.load_data(file_path)
        logger.info(f"成功加载PDF文档，共{len(documents)}个文档片段")
        return documents
    
    def _load_txt(self, file_path: str) -> List[Document]:
        """加载TXT文件"""
        logger.info(f"开始加载TXT文件: {file_path}")
        content = None
        
        # 尝试不同编码读取文件
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        last_error = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                logger.info(f"成功使用 {encoding} 编码读取TXT文件")
                break
            except UnicodeDecodeError as e:
                last_error = e
                logger.debug(f"使用 {encoding} 编码读取失败，尝试下一种编码")
                continue
        
        if content is None:
            error_msg = f"无法读取TXT文件，所有编码尝试均失败: {last_error}"
            logger.error(error_msg)
            raise ValueError(error_msg)
        
        # 创建Document对象
        document = Document(
            text=content,
            metadata={
                'file_name': os.path.basename(file_path),
                'file_type': 'txt',
                'file_size': os.path.getsize(file_path),
                'encoding': encoding  # 记录成功的编码
            }
        )
        
        logger.info(f"成功加载TXT文档，内容长度: {len(content)}")
        return [document]
    
    def _load_word(self, file_path: str) -> List[Document]:
        """加载DOC/DOCX文件"""
        try:
            import docx
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("需要安装python-docx库来处理Word文档: pip install python-docx")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.docx':
            return self._load_docx(file_path)
        elif file_ext == '.doc':
            return self._load_doc(file_path)
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """加载DOCX文件"""
        try:
            import docx
        except ImportError:
            raise ImportError("需要安装python-docx库: pip install python-docx")
        
        doc = docx.Document(file_path)
        
        # 提取所有段落文本
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text.strip())
        
        content = '\n'.join(paragraphs)
        
        # 创建Document对象
        document = Document(
            text=content,
            metadata={
                'file_name': os.path.basename(file_path),
                'file_type': 'docx',
                'file_size': os.path.getsize(file_path),
                'paragraphs_count': len(paragraphs)
            }
        )
        
        logger.info(f"成功加载DOCX文档，段落数: {len(paragraphs)}")
        return [document]
    
    def _load_doc(self, file_path: str) -> List[Document]:
        """加载DOC文件"""
        try:
            import win32com.client
        except ImportError:
            # 如果没有win32com，尝试使用python-docx2txt
            try:
                import docx2txt
                content = docx2txt.process(file_path)
            except ImportError:
                raise ImportError("处理.doc文件需要安装docx2txt库: pip install docx2txt")
        else:
            # 使用win32com处理（仅Windows）
            word = win32com.client.Dispatch("Word.Application")
            word.visible = False
            doc = word.Documents.Open(file_path)
            content = doc.Content.Text
            doc.Close()
            word.Quit()
        
        # 创建Document对象
        document = Document(
            text=content,
            metadata={
                'file_name': os.path.basename(file_path),
                'file_type': 'doc',
                'file_size': os.path.getsize(file_path)
            }
        )
        
        logger.info(f"成功加载DOC文档，内容长度: {len(content)}")
        return [document]
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        提取文件的元数据
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict[str, Any]: 元数据字典
        """
        file_ext = Path(file_path).suffix.lower()
        
        # 确保文件扩展名不带点号
        file_type = file_ext.replace('.', '') if file_ext else 'unknown'
        
        # 基础元数据
        metadata = {
            'file_name': os.path.basename(file_path),
            'file_type': file_type,  # 文件类型（不含点号）
            'file_size': os.path.getsize(file_path),
            'file_path': file_path
        }
        
        try:
            if file_ext == '.pdf':
                metadata.update(self._extract_pdf_metadata(file_path))
            elif file_ext == '.docx':
                metadata.update(self._extract_docx_metadata(file_path))
            # TXT和DOC文件的元数据较少，使用基础元数据即可
        except Exception as e:
            logger.warning(f"提取元数据失败: {str(e)}")
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取PDF元数据"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            metadata = doc.metadata
            
            return {
                'title': metadata.get('title', ''),
                'author': metadata.get('author', ''),
                'subject': metadata.get('subject', ''),
                'creator': metadata.get('creator', ''),
                'producer': metadata.get('producer', ''),
                'creation_date': metadata.get('creationDate', ''),
                'modification_date': metadata.get('modDate', ''),
                'pages': doc.page_count
            }
        except Exception as e:
            logger.error(f"提取PDF元数据失败: {str(e)}")
            return {}
    
    def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取DOCX元数据"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            core_props = doc.core_properties
            
            return {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'creator': core_props.author or '',
                'creation_date': core_props.created.isoformat() if core_props.created else '',
                'modification_date': core_props.modified.isoformat() if core_props.modified else '',
                'paragraphs': len(doc.paragraphs)
            }
        except Exception as e:
            logger.error(f"提取DOCX元数据失败: {str(e)}")
            return {}
    
    def process_document(
        self, 
        document_id: str, 
        storage_type: str, 
        file_path: str, 
        cos_object_key: str = None
    ) -> Dict[str, Any]:
        """
        处理文档并返回处理结果
        
        Args:
            document_id: 文档ID
            storage_type: 存储类型
            file_path: 文件路径
            cos_object_key: COS对象键
            
        Returns:
            处理结果字典
        """
        try:
            # 获取文件内容
            from app.utils.file_storage import file_storage_manager
            
            file_content = file_storage_manager.get_file_content(
                document_id=document_id,
                storage_type=storage_type,
                file_path=file_path,
                cos_object_key=cos_object_key
            )
            
            if not file_content:
                return {
                    "success": False,
                    "error": "无法获取文件内容",
                    "chunks": [],
                    "metadata": {},
                    "chunk_count": 0
                }
            
            # 创建临时文件
            import tempfile
            import os
            from pathlib import Path
            
            # 从文件路径或cos_object_key获取文件扩展名
            if cos_object_key:
                file_ext = Path(cos_object_key).suffix.lower()
            else:
                file_ext = Path(file_path).suffix.lower()
            
            # 确保文件扩展名包含点号且为小写
            if not file_ext:
                # 如果没有扩展名，尝试从文件名推断
                if cos_object_key:
                    logger.warning(f"COS对象键 {cos_object_key} 没有扩展名")
                else:
                    logger.warning(f"文件路径 {file_path} 没有扩展名")
                # 默认为txt，避免后续处理出错
                file_ext = ".txt"
            elif not file_ext.startswith('.'):
                file_ext = '.' + file_ext
            
            file_ext = file_ext.lower()
            
            # 检查文件扩展名是否在支持列表中
            if file_ext not in self.SUPPORTED_TYPES:
                logger.warning(f"文件扩展名 {file_ext} 不在支持列表中: {list(self.SUPPORTED_TYPES.keys())}")
            
            # 记录文件类型信息
            file_type = file_ext.replace('.', '') if file_ext else 'unknown'
            
            logger.info(f"创建临时文件，扩展名: {file_ext}，文件类型: {file_type}")
            
            temp_file_path = None
            try:
                # 确保临时文件有正确的扩展名
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
                    temp_file.write(file_content)
                    temp_file_path = temp_file.name
                
                logger.info(f"临时文件创建成功: {temp_file_path}")
                
                # 验证临时文件是否有正确的扩展名
                temp_file_ext = Path(temp_file_path).suffix.lower()
                logger.info(f"临时文件扩展名: {temp_file_ext}，期望扩展名: {file_ext}")
                
                if temp_file_ext != file_ext:
                    logger.warning(f"临时文件扩展名 {temp_file_ext} 与期望不符 {file_ext}，可能会影响文件类型识别")
                
                # 加载文档
                try:
                    documents = self.load_data(temp_file_path)
                except FileNotFoundError as e:
                    logger.error(f"临时文件不存在: {str(e)}")
                    return {
                        "success": False,
                        "error": f"临时文件不存在: {str(e)}",
                        "chunks": [],
                        "metadata": {},
                        "chunk_count": 0
                    }
                except ValueError as e:
                    if "文件不是有效的PDF格式" in str(e) and file_ext != ".pdf":
                        # 特殊处理：文件扩展名与处理方法不匹配
                        logger.error(f"文件扩展名 {file_ext} 与实际内容不匹配，尝试作为txt处理")
                        
                        # 尝试作为TXT处理
                        try:
                            # 尝试以不同编码读取文本内容
                            content = None
                            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
                            for encoding in encodings:
                                try:
                                    with open(temp_file_path, 'r', encoding=encoding) as f:
                                        content = f.read()
                                    if content:
                                        break
                                except UnicodeDecodeError:
                                    continue
                            
                            if content:
                                # 创建Document对象
                                document = Document(
                                    text=content,
                                    metadata={
                                        'file_name': os.path.basename(file_path or cos_object_key or "unknown"),
                                        'file_type': 'txt',
                                        'file_size': os.path.getsize(temp_file_path),
                                        'encoding': encoding
                                    }
                                )
                                documents = [document]
                                logger.info(f"作为文本文件处理成功，内容长度: {len(content)}")
                            else:
                                # 如果所有编码都失败，返回原错误
                                logger.error(f"作为文本处理也失败")
                                return {
                                    "success": False,
                                    "error": str(e),
                                    "chunks": [],
                                    "metadata": {},
                                    "chunk_count": 0
                                }
                        except Exception as txt_e:
                            logger.error(f"尝试作为文本处理失败: {str(txt_e)}")
                            return {
                                "success": False,
                                "error": f"文件处理失败: {str(e)}，尝试作为文本处理也失败: {str(txt_e)}",
                                "chunks": [],
                                "metadata": {},
                                "chunk_count": 0
                            }
                    else:
                        # 其他值错误，直接返回
                        return {
                            "success": False,
                            "error": str(e),
                            "chunks": [],
                            "metadata": {},
                            "chunk_count": 0
                        }
                
                # 提取元数据
                metadata = self.extract_metadata(temp_file_path)
                
                # 使用SentenceSplitter进行语义分块
                logger.info("开始进行语义分块...")
                text_splitter = SentenceSplitter(
                    chunk_size=1024,
                    chunk_overlap=200
                )
                
                nodes = text_splitter.get_nodes_from_documents(documents)
                logger.info(f"语义分块完成，生成 {len(nodes)} 个节点")
                
                chunks = []
                for i, node in enumerate(nodes):
                    chunk_text = node.get_content()
                    if not chunk_text.strip():
                        continue
                        
                    # 生成UUID格式的chunk_id
                    import uuid
                    chunk_id = str(uuid.uuid4())
                    
                    # 合并节点元数据和文件级元数据
                    chunk_metadata = metadata.copy()
                    chunk_metadata.update(node.metadata)
                    chunk_metadata.update({
                        "chunk_index": i,
                        "document_id": document_id
                    })
                    
                    chunks.append({
                        "content": chunk_text,
                        "chunk_id": chunk_id,
                        "chunk_index": i,
                        "chunk_length": len(chunk_text),
                        "document_id": document_id,
                        "metadata": chunk_metadata
                    })
                
                # 设置页数（对于非PDF文件，设置为1）
                if 'pages' not in metadata:
                    if metadata.get('file_type') == 'pdf':
                        metadata['pages'] = metadata.get('page_count', 1)
                    else:
                        metadata['pages'] = 1
                
                logger.info(f"文档处理成功，生成{len(chunks)}个文本块")
                return {
                    "success": True,
                    "error": None,
                    "chunks": chunks,
                    "metadata": metadata,
                    "chunk_count": len(chunks)
                }
            except FileNotFoundError as e:
                # 文件不存在错误
                logger.error(f"文件不存在: {str(e)}")
                return {
                    "success": False,
                    "error": f"文件不存在: {str(e)}",
                    "chunks": [],
                    "metadata": {},
                    "chunk_count": 0
                }
            except ValueError as e:
                # 值错误（包括文件格式错误）
                logger.error(f"文件格式或值错误: {str(e)}")
                return {
                    "success": False,
                    "error": str(e),
                    "chunks": [],
                    "metadata": {},
                    "chunk_count": 0
                }
            except Exception as e:
                # 其他未预期的错误
                logger.error(f"处理文档时发生未知错误: {str(e)}", exc_info=True)
                return {
                    "success": False,
                    "error": f"处理文档失败: {str(e)}",
                    "chunks": [],
                    "metadata": {},
                    "chunk_count": 0
                }
            finally:
                # 清理临时文件
                try:
                    if temp_file_path and os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)
                        logger.debug(f"已清理临时文件: {temp_file_path}")
                except Exception as e:
                    logger.warning(f"清理临时文件失败: {str(e)}")
        except Exception as e:
            logger.error(f"处理文档过程中发生异常: {str(e)}", exc_info=True)
            return {
                "success": False,
                "error": f"处理文档失败: {str(e)}",
                "chunks": [],
                "metadata": {},
                "chunk_count": 0
            }